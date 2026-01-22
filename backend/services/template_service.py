from docx import Document
from fpdf import FPDF
import tempfile

def gerar_word_com_template(versiculos, output_pdf_path):
    """
    Recebe a lista de versículos e gera:
    1. Word com tabela 3x4 usando os versículos
    2. PDF final
    Retorna o caminho do PDF.
    """

    # 1️⃣ Cria documento Word temporário
    doc = Document()
    doc.add_heading("Versículos Selecionados", level=1)

    # Tabela 3x4
    tabela = doc.add_table(rows=3, cols=4)
    tabela.style = 'Table Grid'

    i = 0
    for row in tabela.rows:
        for cell in row.cells:
            if i < len(versiculos):
                v = versiculos[i]
                cell.text = f"{v['texto']}\n({v['endereco']})"
                i += 1
            else:
                cell.text = ""

    # Salva Word temporário
    tmp_word = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_word.name)

    # 2️⃣ Converte Word para PDF simples com FPDF
    # OBS: FPDF não lê docx, então vamos simplificar: PDF com texto dos versículos
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, "Versículos Selecionados", ln=True, align="C")
    pdf.ln(5)

    for v in versiculos:
        pdf.multi_cell(0, 8, f"{v['texto']} ({v['endereco']})")
        pdf.ln(1)

    pdf.output(output_pdf_path)
    return output_pdf_path
